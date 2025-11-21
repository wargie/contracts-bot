import os
import textwrap
from typing import Dict

from jinja2 import Environment, BaseLoader
from docx import Document

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab import rl_config


DEFAULT_TEMPLATE = """
ДОГОВОР ОКАЗАНИЯ УСЛУГ № {{ number }}

г. {{ city }} «{{ date }}»

{{ customer.name }}, ИНН {{ customer.inn }}{% if customer.kpp %} КПП {{ customer.kpp }}{% endif %}, ОГРН {{ customer.ogrn or '-' }}, адрес: {{ customer.address or '-' }}, в лице {{ customer.manager or 'представителя' }}, именуем{{ 'ая' if 'ИП' in customer.name else 'ый' }} далее «Заказчик», с одной стороны, и {{ contractor.name }}, ИНН {{ contractor.inn }}{% if contractor.kpp %} КПП {{ contractor.kpp }}{% endif %}, ОГРН {{ contractor.ogrn or '-' }}, адрес: {{ contractor.address or '-' }}, в лице {{ contractor.manager or 'представителя' }}, именуемый далее «Исполнитель», совместно — «Стороны», заключили настоящий договор о нижеследующем:

1. Предмет договора
1.1. Исполнитель обязуется оказать услуги: {{ subject }}.

2. Стоимость и порядок расчетов
2.1. Стоимость услуг составляет {{ price }}.
2.2. Порядок расчетов: {{ payment_terms }}.

3. Сроки
3.1. Срок исполнения: {{ term }}.

4. Ответственность Сторон
4.1. Стороны несут ответственность в соответствии с действующим законодательством {{ jurisdiction }}.{% if penalties %}
4.2. Штрафные санкции: {{ penalties }}.{% endif %}

5. Форс-мажор
5.1. Стороны освобождаются от ответственности за частичное или полное неисполнение обязательств, вызванное обстоятельствами непреодолимой силы.

6. Реквизиты и подписи Сторон

Заказчик:
{{ customer.name }}
ИНН {{ customer.inn }}{% if customer.kpp %} КПП {{ customer.kpp }}{% endif %}
ОГРН {{ customer.ogrn or '-' }}
Адрес: {{ customer.address or '-' }}
Банк: {{ customer.bank_name or '-' }}, БИК {{ customer.bank_bik or '-' }}
Р/с {{ customer.bank_account or '-' }}, к/с {{ customer.bank_corr or '-' }}

Исполнитель:
{{ contractor.name }}
ИНН {{ contractor.inn }}{% if contractor.kpp %} КПП {{ contractor.kpp }}{% endif %}
ОГРН {{ contractor.ogrn or '-' }}
Адрес: {{ contractor.address or '-' }}
Банк: {{ contractor.bank_name or '-' }}, БИК {{ contractor.bank_bik or '-' }}
Р/с {{ contractor.bank_account or '-' }}, к/с {{ contractor.bank_corr or '-' }}

Подписи:
______________/____________________/     _____________/____________________/

(Отметка: реквизиты проверены: {{ verification_status }})
"""


def render_text(context: Dict) -> str:
    env = Environment(loader=BaseLoader(), autoescape=False, trim_blocks=True, lstrip_blocks=True)
    template = env.from_string(DEFAULT_TEMPLATE)
    return template.render(**context)


def text_to_docx(text: str, out_path: str) -> str:
    doc = Document()
    for para in text.split("\n\n"):
        p = doc.add_paragraph()
        for line in para.split("\n"):
            p.add_run(line)
            p.add_run("\n")
    doc.save(out_path)
    return out_path


# ---------------- PDF ----------------

def _extend_ttf_search_path():
    """Добавляем каталоги, где искать TTF (включая app/fonts)."""
    add = [
        os.path.abspath(os.path.join("app", "assets", "fonts")),
        os.path.abspath(os.path.join("app", "fonts")),          # <-- ДОБАВЛЕНО
        os.path.abspath(os.path.join("assets", "fonts")),
        os.path.abspath("fonts"),
        r"C:\Windows\Fonts",
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts/truetype",
        "/Library/Fonts",
        "/System/Library/Fonts",
    ]
    for p in add:
        if p and p not in rl_config.TTFSearchPath:
            rl_config.TTFSearchPath.append(p)


def _register_font(name: str, path: str) -> bool:
    try:
        pdfmetrics.registerFont(TTFont(name, path))
        return True
    except Exception:
        return False


def _ensure_cyrillic_font() -> str:
    """
    Ищем и регистрируем шрифт с кириллицей.
    Возвращаем имя зарегистрированного шрифта.
    """
    _extend_ttf_search_path()

    candidates = [
        ("DejaVuSans", ["DejaVuSans.ttf", "dejavusans.ttf"]),
        ("Arial", ["Arial.ttf", "arial.ttf"]),
        ("LiberationSans", ["LiberationSans-Regular.ttf", "LiberationSans.ttf"]),
        ("TimesNewRoman", ["Times New Roman.ttf", "times.ttf", "Times.ttf"]),
    ]

    for font_name, files in candidates:
        for base in rl_config.TTFSearchPath:
            for fn in files:
                fpath = os.path.join(base, fn)
                if os.path.exists(fpath) and _register_font(font_name, fpath):
                    return font_name

    # запасные абсолютные пути
    absolutes = [
        r"C:\Windows\Fonts\arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for fpath in absolutes:
        if os.path.exists(fpath) and _register_font("ArialSys", fpath):
            return "ArialSys"

    return "Helvetica"  # если дойдём сюда — не будет кириллицы


def _wrap_by_width(c: canvas.Canvas, text: str, max_width: float) -> list[str]:
    """Перенос строк по ширине с учётом метрик шрифта."""
    words = text.split(" ")
    lines, cur = [], ""
    for w in words:
        nxt = (cur + " " + w) if cur else w
        if c.stringWidth(nxt) <= max_width:
            cur = nxt
        else:
            if cur:
                lines.append(cur)
            if c.stringWidth(w) > max_width:
                chunk = ""
                for ch in w:
                    if c.stringWidth(chunk + ch) <= max_width:
                        chunk += ch
                    else:
                        lines.append(chunk)
                        chunk = ch
                cur = chunk
            else:
                cur = w
    if cur:
        lines.append(cur)
    return lines


def text_to_pdf(text: str, out_path: str) -> str:
    """Генерация PDF с кириллицей из простого текста."""
    font_name = _ensure_cyrillic_font()

    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4

    font_size = 11
    line_height = 14
    c.setFont(font_name, font_size)

    left = 40
    right = 40
    top = 40
    bottom = 40
    y = height - top
    max_width = width - left - right

    for paragraph in text.split("\n\n"):
        for line in paragraph.split("\n"):
            if " " not in line and c.stringWidth(line) > max_width:
                for wrapped in textwrap.wrap(line, width=100):
                    for l2 in _wrap_by_width(c, wrapped, max_width):
                        c.drawString(left, y, l2)
                        y -= line_height
                        if y < bottom:
                            c.showPage()
                            c.setFont(font_name, font_size)
                            y = height - top
            else:
                for l2 in _wrap_by_width(c, line, max_width):
                    c.drawString(left, y, l2)
                    y -= line_height
                    if y < bottom:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = height - top
        y -= line_height
        if y < bottom:
            c.showPage()
            c.setFont(font_name, font_size)
            y = height - top

    c.save()
    return out_path